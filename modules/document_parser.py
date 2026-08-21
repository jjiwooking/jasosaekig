from __future__ import annotations

from io import BytesIO
from pathlib import Path


def extract_text(uploaded_file) -> str:
    name = (uploaded_file.name or "").lower()
    raw = uploaded_file.getvalue()

    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(raw))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()

    if name.endswith(".docx"):
        from docx import Document

        doc = Document(BytesIO(raw))
        chunks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(chunks).strip()

    if name.endswith(".txt") or name.endswith(".md"):
        for encoding in ("utf-8-sig", "utf-8", "cp949"):
            try:
                return raw.decode(encoding).strip()
            except UnicodeDecodeError:
                continue

    raise ValueError(f"지원하지 않는 파일 형식입니다: {Path(name).suffix}")
